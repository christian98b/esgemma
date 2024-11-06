import re
from typing import List, Tuple
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.metrics import accuracy_score, precision_score, recall_score, f1_score, confusion_matrix
from collections import Counter
from docx import Document
from docx.shared import Inches
import os
import xml.etree.ElementTree as ET


"""
Naming convention:
https://stackoverflow.com/questions/46663013/what-is-y-true-and-y-pred-when-creating-a-custom-metric-in-keras

These functions work as follows

always call evaluate_llm_metrics first

With the results call save_to_docx or print_results

The user should call the evaluate_llm_output function that should get a List of Tuples that represent the ground truth [(Target,Year,Quote)]
"""

def __parse_xml_string(xml_string: str) -> Tuple[str, str, str]:
    target_match = re.search(r'<end_target>(.*?)</end_target>', xml_string, re.DOTALL)
    year_match = re.search(r'<end_target_year>(.*?)</end_target_year>', xml_string, re.DOTALL)
    quote_match = re.search(r'<quote>(.*?)</quote>', xml_string, re.DOTALL)
    
    target = target_match.group(1) if target_match else "N/A"
    
    #Used for gemma if wrongly formatted
    #if(target_match == None):
    #    target_match = re.search(r'<classification>(.*?)</classification>', xml_string, re.DOTALL)
    #    target = target_match.group(1) if target_match else "N/A"

    year = year_match.group(1) if year_match else "N/A"
    quote = quote_match.group(1) if quote_match else "N/A"
    
    return target, year, quote

def __calculate_metrics(y_true: List[str], y_pred: List[str]) -> dict:
    accuracy = accuracy_score(y_true, y_pred)
    precision = precision_score(y_true, y_pred, average='weighted', zero_division=0)
    recall = recall_score(y_true, y_pred, average='weighted', zero_division=0)
    f1 = f1_score(y_true, y_pred, average='weighted', zero_division=0)
    
    #Create the classes as a set. For example is ['Carbon neutral(ity)', 'Emissions reduction target', 'Net zero', 'No target']
    #These classes can contain N/A if the model cant extract a value based on the pars
    classes = sorted(set(y_true + y_pred))
    class_metrics = {}
    
    #Iterate over classes compare with the values in y,true and predicted. 1 if the same, 0 if different.
    for cls in classes:
        cls_y_true = [1 if y == cls else 0 for y in y_true] #Is a list that looks for example like this. [1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 0, 0, 0, 0, 0, 0, 0, 0, 0, 1, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0]
        cls_y_pred = [1 if y == cls else 0 for y in y_pred] #Is the same but with the true 0 and 1s
        ##This essentially creates two 1, and 0 lists that can be compared to find out if the right class is predicted
        
        #Accuracy will not be predicited as its not a good metric for individual classes. Not using weighted as we dont need to look on class imbalances as its only calculating for a single class.
        cls_precision = precision_score(cls_y_true, cls_y_pred, zero_division=0)
        cls_recall = recall_score(cls_y_true, cls_y_pred, zero_division=0)
        cls_f1 = f1_score(cls_y_true, cls_y_pred, zero_division=0)
        
        class_metrics[cls] = {
            "precision": cls_precision,
            "recall": cls_recall,
            "f1": cls_f1
        }
    
    return {
        "accuracy": accuracy,
        "precision": precision,
        "recall": recall,
        "f1": f1,
        "class_metrics": class_metrics
    }

def __plot_confusion_matrix(cm : np.ndarray, classes : list[str], title:str):
    #https://seaborn.pydata.org/generated/seaborn.heatmap.html
    plt.figure(figsize=(10, 8))
    sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', xticklabels=classes, yticklabels=classes)
    plt.title(title)
    plt.ylabel('True label')
    plt.xlabel('Predicted label')
    plt.tight_layout()
    plt.show()


def __plot_confusion_matrix_and_save(cm : np.ndarray, classes : list[str], title : str, filename : str):
    plt.figure(figsize=(10, 8))
    sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', xticklabels=classes, yticklabels=classes) #Display as decimal
    plt.title(title)
    plt.ylabel('True label')
    plt.xlabel('Predicted label')
    plt.tight_layout()
    plt.savefig(filename)
    plt.close()  # Close the plot to free up memory


def  __is_valid_xml(xml_string):
    target_match = re.search(r'<answer>(.*?)</answer>', xml_string, re.DOTALL)
    return True if target_match else False

def evaluate_llm_output(ground_truth: List[Tuple[str, str, str]], predictions: List[str]) -> dict:
    """
    Evaluate the performance of an LLM output against ground truth data.

    This function initializes the evaluation process, returning various performance metrics 
    that can be printed or saved using `print_results` or `save_results_to_docx`.

    Parameters
    ----------
    ground_truth : List[Tuple[str, str, str]]
        A list of ground truth tuples in the form (target, year, quote).

    predictions : List[str]
        A list of LLM-generated XML strings or similar formatted strings. 
        The output can contain other things besides the XML string as long as it includes it.

    Returns
    -------
    dict
        A dictionary containing various evaluation metrics and results, structured as follows:
        
        {
            "target_metrics": {
                "accuracy": float,
                "precision": float,
                "recall": float,
                "f1": float,
                "class_metrics": {
                    "class1": {
                        "precision": float,
                        "recall": float,
                        "f1": float
                    },
                    "class2": {
                        "precision": float,
                        "recall": float,
                        "f1": float
                    },
                    ...
                }
            },
            "year_metrics": {
                "accuracy": float,
                "precision": float,
                "recall": float,
                "f1": float,
                "class_metrics": {
                    "year1": {
                        "precision": float,
                        "recall": float,
                        "f1": float
                    },
                    "year2": {
                        "precision": float,
                        "recall": float,
                        "f1": float
                    },
                    ...
                }
            },
            "overall_accuracy": float,
            "valid_xml_percentage": float,
            "quote_accuracy": float,
            "target_cm": numpy.ndarray,
            "target_classes": List[str],
            "year_cm": numpy.ndarray,
            "year_classes": List[str]
        }

    Examples
    --------
    >>> ground_truth = [
    ...     ("Carbon neutral(ity)", "2040", "Thus our company aims to achieve carbon neutrality by 2040."),
    ...     ("Sustainability", "2030", "Our goal is to be fully sustainable by 2030.")
    ... ]
    >>> predictions = [
    ...     "<response><target>Carbon neutral(ity)</target><year>2040</year><quote>Thus our company aims to achieve carbon neutrality by 2040.</quote></response>",
    ...     "<response><target>No target</target><year>No target</year><quote>None</quote></response>"
    ... ]
    >>> results = evaluate_llm_output(ground_truth, predictions)
    """

    ### First parse the predictions

    # Parse predictions
    parsed_predictions : list[Tuple[str, str, str]] = [__parse_xml_string(pred) for pred in predictions]
    
    # Prepare data for sklearn metrics
    # ("Carbon neutral(ity)", "2040", "Thus Our company aims to achieve carbon neutrality by 2040. This ambitious goal reflects our commitment to sustainability and reducing our environmental impact.")
    y_true_target : list[str] = [gt[0] for gt in ground_truth]
    y_true_year : list[str] = [gt[1] for gt in ground_truth]
    y_pred_target : list[str] = [pred[0] for pred in parsed_predictions]
    y_pred_year : list[str] = [pred[1] for pred in parsed_predictions]
    

    ##Second calculate metrics


    # Calculate metrics for targets and years
    target_metrics : dict = __calculate_metrics(y_true_target, y_pred_target)
    year_metrics : dict = __calculate_metrics(y_true_year, y_pred_year)
    
    # Calculate overall accuracy (both target and year correct)
    overall_accuracy : float = np.mean([gt[:2] == pred[:2] for gt, pred in zip(ground_truth, parsed_predictions)])
    
    # Calculate percentage of valid XML responses
    # A xml is valid as soon the output tags are detected. Careful could result in a wrong value.
    valid_xml_percentage = sum(1 for pred in predictions if __is_valid_xml(pred)) / len(predictions)
    
    # Calculate quote accuracy
    # A quote is valid if the model either classifies as No target or if extraction and classification is not No target and the quote is in the ground_truth
    quote_accuracy = sum(1 for gt, pred in zip(ground_truth, parsed_predictions) 
                          if (pred[0] == "No target" and pred[1] == "No target" and pred[2] == "None") or 
                             (pred[0] != "No target" and pred[1] != "No target" and pred[2] in gt[2])) / len(ground_truth)
    
    
    
    
    # Third create the confision matrices


    # First create a list with strings that contains all the prediction and ground truth labels for the matrix
    target_classes = sorted(set(y_true_target + y_pred_target))
    year_classes = sorted(set(y_true_year + y_pred_year))
    # Second call the confusion matrix method
    target_cm : np.ndarray = confusion_matrix(y_true_target, y_pred_target, labels=target_classes)
    year_cm : np.ndarray = confusion_matrix(y_true_year, y_pred_year, labels=year_classes)

    return {
        "target_metrics": target_metrics,
        "year_metrics": year_metrics,
        "overall_accuracy": overall_accuracy,
        "valid_xml_percentage": valid_xml_percentage,
        "quote_accuracy": quote_accuracy,
        "target_cm" : target_cm,
        "target_classes" : target_classes,
        "year_cm" : year_cm,
        "year_classes": year_classes
    }



def save_results_to_docx(results: dict, folder_path: str, file_name: str):

    # Ensure the directory exists
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)

    # Create a new Document
    doc = Document()

    # Write Target Metrics
    doc.add_heading('Target Metrics', level=1)
    doc.add_paragraph(f"Overall - Accuracy: {results['target_metrics']['accuracy']:.4f}, "
                      f"Precision: {results['target_metrics']['precision']:.4f}, "
                      f"Recall: {results['target_metrics']['recall']:.4f}, "
                      f"F1: {results['target_metrics']['f1']:.4f}")

    doc.add_paragraph('Class-specific metrics:')
    for cls, metrics in results['target_metrics']['class_metrics'].items():
        doc.add_paragraph(f"  {cls} - Precision: {metrics['precision']:.4f}, "
                          f"Recall: {metrics['recall']:.4f}, "
                          f"F1: {metrics['f1']:.4f}")

    # Write Year Metrics
    doc.add_heading('Year Metrics', level=1)
    doc.add_paragraph(f"Overall - Accuracy: {results['year_metrics']['accuracy']:.4f}, "
                      f"Precision: {results['year_metrics']['precision']:.4f}, "
                      f"Recall: {results['year_metrics']['recall']:.4f}, "
                      f"F1: {results['year_metrics']['f1']:.4f}")

    doc.add_paragraph('Class-specific metrics:')
    for cls, metrics in results['year_metrics']['class_metrics'].items():
        doc.add_paragraph(f"  {cls} - Precision: {metrics['precision']:.4f}, "
                          f"Recall: {metrics['recall']:.4f}, "
                          f"F1: {metrics['f1']:.4f}")

    # Write Quote Accuracy and Overall Accuracy
    doc.add_heading('Other Metrics', level=1)
    doc.add_paragraph(f"Quote Accuracy: {results['quote_accuracy']:.4f}")
    doc.add_paragraph(f"Overall Accuracy: {results['overall_accuracy']:.4f}")
    doc.add_paragraph(f"Valid XML Percentage: {results['valid_xml_percentage']:.4f}")

    # Save and insert confusion matrices as images
    target_cm_filename = os.path.join(folder_path, "target_confusion_matrix.png")
    year_cm_filename = os.path.join(folder_path, "year_confusion_matrix.png")

    # Plot and save confusion matrices
    __plot_confusion_matrix_and_save(results['target_cm'], results['target_classes'], 
                                   "Target Confusion Matrix", target_cm_filename)
    __plot_confusion_matrix_and_save(results['year_cm'], results['year_classes'], 
                                   "Year Confusion Matrix", year_cm_filename)

    # Insert confusion matrix images into the document
    doc.add_heading('Confusion Matrices', level=1)

    doc.add_paragraph('Target Confusion Matrix:')
    doc.add_picture(target_cm_filename, width=Inches(6))  # Adjust the width to fit the document

    doc.add_paragraph('Year Confusion Matrix:')
    doc.add_picture(year_cm_filename, width=Inches(6))

    # Save the document
    file_path = os.path.join(folder_path, f"{file_name}.docx")
    doc.save(file_path)

    # Remove the temporary confusion matrix images
    os.remove(target_cm_filename)
    os.remove(year_cm_filename)

    print(f"Results saved to {file_path}")



def print_results(results : dict):
    print("Target Metrics:")
    print(f"Overall - Accuracy: {results['target_metrics']['accuracy']:.4f}, Precision: {results['target_metrics']['precision']:.4f}, Recall: {results['target_metrics']['recall']:.4f}, F1: {results['target_metrics']['f1']:.4f}")
    print("Class-specific metrics:")
    for cls, metrics in results['target_metrics']['class_metrics'].items():
        print(f"  {cls} - Precision: {metrics['precision']:.4f}, Recall: {metrics['recall']:.4f}, F1: {metrics['f1']:.4f}")
    __plot_confusion_matrix(results['target_cm'], results['target_classes'], "Target Confusion Matrix")

    print("\nYear Metrics:")
    print(f"Overall - Accuracy: {results['year_metrics']['accuracy']:.4f}, Precision: {results['year_metrics']['precision']:.4f}, Recall: {results['year_metrics']['recall']:.4f}, F1: {results['year_metrics']['f1']:.4f}")
    print("Class-specific metrics:")
    for cls, metrics in results['year_metrics']['class_metrics'].items():
        print(f"  {cls} - Precision: {metrics['precision']:.4f}, Recall: {metrics['recall']:.4f}, F1: {metrics['f1']:.4f}")
    __plot_confusion_matrix(results['year_cm'], results['year_classes'], "Year Confusion Matrix")

    print(f"\nQuote Accuracy: {results['quote_accuracy']:.4f}")
    print(f"\nOverall Accuracy: {results['overall_accuracy']:.4f}")
    print(f"Valid XML Percentage: {results['valid_xml_percentage']:.4f}")